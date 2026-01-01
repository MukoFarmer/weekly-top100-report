# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# def write_weekly_report(result: dict, output_path: str):
#     doc = Document()
#     # =========================
#     # TITLE
#     # =========================
#     title = doc.add_heading(
#         f"Weekly Performance Summary – Week {result['week']}",
#         level=1
#     )
#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     doc.add_paragraph(
#         f"Comparison: Week {result['week']} vs Week {result['previous_week']}\n"
#     )
#     # =========================
#     # CONTRIBUTORS / DETRACTORS
#     # =========================
#     def add_section(title_text, data):
#         doc.add_heading(title_text, level=2)
#         for group, items in data.items():
#             doc.add_heading(group.upper(), level=3)
#             if not items:
#                 doc.add_paragraph("N/A")
#                 continue
#             table = doc.add_table(rows=1, cols=4)
#             hdr = table.rows[0].cells
#             hdr[0].text = "Seller"
#             hdr[1].text = f"GMS {result['week']}"
#             hdr[2].text = f"GMS {result['previous_week']}"
#             hdr[3].text = "WoW Diff"
#             for r in items:
#                 row = table.add_row().cells
#                 row[0].text = r["sp_name"]
#                 row[1].text = f"{r[f'gms_{result['week']}']:,}"
#                 row[2].text = f"{r[f'gms_{result['previous_week']}']:,}"
#                 row[3].text = f"{r['diff']:,}"
#     add_section("Top Contributors", result["contributors"])
#     add_section("Top Detractors", result["detractors"])
#     # =========================
#     # PARITY REPORT
#     # =========================
#     doc.add_heading("Selection Parity Analysis", level=2)
#     for line in result["parity_report_text"].split("\n"):
#         p = doc.add_paragraph(line)
#         p.paragraph_format.space_after = Pt(6)
#     # =========================
#     # SAVE
#     # =========================
#     doc.save(output_path)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill_hex: str):
    """Example fill_hex: '000000'."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    font_name: str = "Aptos Narrow",
    font_size_pt: int = 11,
    align: WD_ALIGN_PARAGRAPH | None = None,
):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text if text is not None else "")
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_blank_line(doc: Document, n: int = 1):
    for _ in range(n):
        doc.add_paragraph("")


def _fmt_number(x):
    if x is None:
        return "N/A"
    # Keep integers as integers, floats as 2dp (similar to your sample having decimals)
    if isinstance(x, (int,)) or (isinstance(x, float) and x.is_integer()):
        return f"{int(x):,}"
    if isinstance(x, (float,)):
        return f"{x:,.2f}"
    return str(x)


def _get_gms_value(row: dict, week_key: str):
    """
    Tries common key variants safely:
      - gms_52 style
      - gms_2025W52 style (if you have it)
      - numeric week stored under nested fields
    """
    return row.get(week_key)


def _build_wow_table(doc: Document, groups: dict, week: str, previous_week: str, *, group_title_prefix: str):
    """
    Creates ONE table per section, with:
      Row: merged group title
      Row: black header
      Rows: data
      Repeat for each group
    """
    week_key = f"gms_{week}"
    prev_key = f"gms_{previous_week}"

    table = doc.add_table(rows=0, cols=4)
    table.style = "Normal Table"

    group_order = list(groups.keys())  # preserves insertion order from your input dict

    for gi, group_name in enumerate(group_order):
        items = groups.get(group_name, [])

        # Group title row (merged across 4 columns)
        gr = table.add_row().cells
        gr[0].merge(gr[1]).merge(gr[2]).merge(gr[3])
        _set_cell_text(
            gr[0],
            f"{group_title_prefix} ({group_name})",
            bold=True,
            color=RGBColor(0x00, 0x00, 0x00),
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        # Header row (black shading, white bold text)
        hdr = table.add_row().cells
        headers = ["SP", week_key, prev_key, "gms_difference"]
        for i, h in enumerate(headers):
            _set_cell_shading(hdr[i], "000000")
            _set_cell_text(
                hdr[i],
                h,
                bold=True,
                color=RGBColor(0xFF, 0xFF, 0xFF),
                align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT,
            )

        # Data rows
        if not items:
            r = table.add_row().cells
            r[0].merge(r[1]).merge(r[2]).merge(r[3])
            _set_cell_text(r[0], "N/A", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            continue

        for it in items:
            sp = it.get("sp_name") or it.get("sp") or "N/A"
            gms_w = _get_gms_value(it, week_key)
            gms_p = _get_gms_value(it, prev_key)
            diff = it.get("diff") if "diff" in it else (None if gms_w is None or gms_p is None else (gms_w - gms_p))

            row = table.add_row().cells
            _set_cell_text(row[0], str(sp), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(row[1], _fmt_number(gms_w), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(row[2], _fmt_number(gms_p), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(row[3], _fmt_number(diff), align=WD_ALIGN_PARAGRAPH.RIGHT)

    return table


def write_weekly_report(result: dict, output_path: str):
    """
    Expected inputs (compatible with your current structure, with one adjustment):
      result["contributors"] = { "SAS": [...], "Non-SAS": [...] }  (or similar group keys)
      result["detractors"]   = { "SAS": [...], "Non-SAS": [...] }
      result["parity_report_text"] can still exist; additionally, you may pass:
        result["from_zero_selection_text"] (string, default "N/A.")
        result["wow_parity_increase_text"] (string, default "N/A.")
        result["wow_parity_decrease"]      (dict like {seller: "-92%", ...} or list of {"seller":..., "pct":...})
    """
    doc = Document()

    week = str(result["week"])
    prev = str(result["previous_week"])

    # Intro (matches the shared doc)
    doc.add_paragraph("Dear Zeliha and OHL Team,")
    _add_blank_line(doc, 1)
    doc.add_paragraph(
        "Please find the WoW highlights of our Top 100 SPs below and the details attached and let me know if any further information is needed."
    )
    _add_blank_line(doc, 1)

    # Contributors label + combined table
    doc.add_paragraph("Wow Top GMS Contributors:")
    _add_blank_line(doc, 1)
    _build_wow_table(
        doc,
        groups=result.get("contributors", {}),
        week=week,
        previous_week=prev,
        group_title_prefix="Top Contributors",
    )
    _add_blank_line(doc, 1)

    # Detractors label + combined table
    doc.add_paragraph("Wow Top GMS detractors:")
    _add_blank_line(doc, 1)
    _build_wow_table(
        doc,
        groups=result.get("detractors", {}),
        week=week,
        previous_week=prev,
        group_title_prefix="Top Detractors",
    )
    _add_blank_line(doc, 2)

    # Parity / additional notes (matches your shared doc ordering)
    doc.add_paragraph("From Zero Selection vice versa")
    _add_blank_line(doc, 1)
    doc.add_paragraph(result.get("from_zero_selection_text", "N/A."))
    _add_blank_line(doc, 1)

    doc.add_paragraph("WoW parity increase:")
    _add_blank_line(doc, 1)
    doc.add_paragraph(result.get("wow_parity_increase_text", "N/A."))
    _add_blank_line(doc, 1)

    doc.add_paragraph("WoW parity decrease:")
    _add_blank_line(doc, 1)

    wow_dec = result.get("wow_parity_decrease")
    if not wow_dec:
        doc.add_paragraph("N/A.")
    else:
        # Build the 2-column table like your sample (no header)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Normal Table"

        if isinstance(wow_dec, dict):
            items = list(wow_dec.items())
        else:
            # list[{"seller":..., "pct":...}] or list[("seller","-92%")]
            items = []
            for x in wow_dec:
                if isinstance(x, dict):
                    items.append((x.get("seller") or x.get("sp_name") or x.get("sp") or "N/A", x.get("pct") or x.get("percent") or "N/A"))
                elif isinstance(x, (list, tuple)) and len(x) >= 2:
                    items.append((x[0], x[1]))
                else:
                    items.append((str(x), ""))
        for seller, pct in items:
            r = t.add_row().cells
            _set_cell_text(r[0], str(seller), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(r[1], str(pct), align=WD_ALIGN_PARAGRAPH.LEFT)

    _add_blank_line(doc, 2)
    doc.add_paragraph("Regards.")
    _add_blank_line(doc, 1)

    doc.save(output_path)
